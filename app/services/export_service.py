"""
app/services/export_service.py
FinaCES V1.2 — Word / PDF Export (Async Migration Sprint 2B, GAP-08)

All document generation (docx, weasyprint) is CPU-bound/I/O-bound.
It is isolated from the FastAPI event loop via asyncio.to_thread().
"""

import asyncio
import os
import re
import uuid
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.db.models import MCCGradeReport
from app.services.audit_service import log_event

logger = logging.getLogger(__name__)

# ── Export Directories ────────────────────────────────────────
EXPORT_DIR = Path(os.getenv("FINACES_EXPORT_DIR", "exports"))
EXPORT_DIR.mkdir(parents=True, exist_ok=True)

# ── Section Labels ────────────────────────────────────────────
SECTION_TITLES = {
    "section_01_info":           "1. General Information",
    "section_02_objective":      "2. Purpose of the Report",
    "section_03_scope":          "3. Scope of Analysis",
    "section_04_executive_summary":"4. Executive Summary",
    "section_05_profile":        "5. Bidder Profile",
    "section_06_analysis":       "6. Financial Analysis",
    "section_07_capacity":       "7. Contract Capacity",
    "section_08_red_flags":      "8. Identified Red Flags",
    "section_09_mitigants":      "9. Mitigating Factors",
    "section_10_scoring":        "10. MCC-grade Scoring",
    "section_11_assessment":     "11. Overall Risk Assessment",
    "section_12_recommendation": "12. Recommendation",
    "section_13_limitations":    "13. Analysis Limitations",
    "section_14_conclusion":     "14. Conclusion",
}

RISK_COLORS_HEX = {
    "LOW":      "#27AE60",
    "MEDIUM":   "#F39C12",
    "HIGH":     "#E67E22",
    "CRITICAL": "#E74C3C",
}


# ════════════════════════════════════════════════════════════════
# EXPORT WORD (Async facade)
# ════════════════════════════════════════════════════════════════

async def export_to_word(
    report:  dict,
    case_id: str,
    db:      AsyncSession,
) -> str:
    """
    Generates a Word (.docx) file for the MCC-grade report.

    Generation is delegated via asyncio.to_thread() to avoid blocking
    the event loop. Updates MCCGradeReport.export_word_path in the DB.
    Returns the absolute path of the generated file.
    """
    subdir = EXPORT_DIR / case_id
    subdir.mkdir(parents=True, exist_ok=True)

    ref = _sanitize_filename(
        report.get("section_01_info", "")[:30] or case_id[:8]
    )
    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    file_path = subdir / f"Note_MCC_{ref}_{ts}.docx"

    # ─ Document generation (CPU-bound) → isolated thread ─────
    await asyncio.to_thread(_sync_generate_word, report, str(file_path))

    # ─ ORM Update ─────────────────────────────────────────────
    await _persist_export_path(
        report_id=report.get("report_id"),
        word_path=str(file_path),
        db=db,
    )

    # ─ Audit ─────────────────────────────────────────────────
    await log_event(
        db=db,
        event_type="REPORT_EXPORTED",
        entity_type="MCCGradeReport",
        entity_id=report.get("report_id"),
        case_id=case_id,
        description=f"Word Export: {file_path.name}",
        new_value={"format": "DOCX", "path": str(file_path)},
    )

    logger.info(f"Word export generated: {file_path}")
    return str(file_path)


# ════════════════════════════════════════════════════════════════
# EXPORT PDF (Async facade)
# ════════════════════════════════════════════════════════════════

async def export_to_pdf(
    report:  dict,
    case_id: str,
    db:      AsyncSession,
) -> str:
    """
    Generates a PDF file via WeasyPrint (HTML fallback if unavailable).

    Generation is delegated via asyncio.to_thread() to avoid blocking
    the event loop. Updates MCCGradeReport.export_pdf_path in the DB.
    Returns the absolute path of the generated file.
    """
    subdir = EXPORT_DIR / case_id
    subdir.mkdir(parents=True, exist_ok=True)

    ref = _sanitize_filename(case_id[:8])
    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    pdf_path = subdir / f"Note_MCC_{ref}_{ts}.pdf"

    # ─ Document generation (CPU-bound) → isolated thread ─────
    actual_path = await asyncio.to_thread(_sync_generate_pdf, report, str(pdf_path), str(subdir), ref, ts)

    # ─ ORM Update ─────────────────────────────────────────────
    await _persist_export_path(
        report_id=report.get("report_id"),
        pdf_path=actual_path,
        db=db,
    )

    # ─ Audit ─────────────────────────────────────────────────
    await log_event(
        db=db,
        event_type="REPORT_EXPORTED",
        entity_type="MCCGradeReport",
        entity_id=report.get("report_id"),
        case_id=case_id,
        description=f"PDF Export: {Path(actual_path).name}",
        new_value={"format": "PDF", "path": actual_path},
    )

    logger.info(f"PDF export generated: {actual_path}")
    return actual_path


# ════════════════════════════════════════════════════════════════
# HELPERS ASYNC DB
# ════════════════════════════════════════════════════════════════

async def _persist_export_path(
    report_id: Optional[str],
    db:        AsyncSession,
    word_path: Optional[str] = None,
    pdf_path:  Optional[str] = None,
) -> None:
    """Updates export paths in the MCCGradeReport ORM."""
    if not report_id:
        return
    result = await db.execute(
        select(MCCGradeReport).where(MCCGradeReport.id == uuid.UUID(report_id))
    )
    report_orm = result.scalars().first()
    if report_orm:
        if word_path:
            report_orm.export_word_path = word_path
        if pdf_path:
            report_orm.export_pdf_path = pdf_path
        report_orm.updated_at = datetime.now(timezone.utc)
        await db.commit()


# ════════════════════════════════════════════════════════════════
# SYNCHRONOUS GENERATION — Word (executed via asyncio.to_thread)
# ════════════════════════════════════════════════════════════════

def _sync_generate_word(report: dict, file_path: str) -> None:
    """Generates the .docx file synchronously (called via to_thread)."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    _configure_document(doc)
    _add_cover_page(doc, report)
    _add_page_break(doc)
    _add_table_of_contents(doc)
    _add_page_break(doc)
    _add_all_sections(doc, report)
    _add_footer(doc, report)
    doc.save(file_path)


def _configure_document(doc) -> None:
    from docx.shared import Pt, Cm
    from docx.shared import RGBColor

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1A, 0x52, 0x7A)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)


def _add_cover_page(doc, report: dict) -> None:
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FINANCIAL ANALYSIS REPORT")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x52, 0x7A)
    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("MCC/MCA Standard — Fiduciary Grade")
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    doc.add_paragraph("\n\n")

    risk_class = _extract_risk_class(report)
    if risk_class:
        risk_para = doc.add_paragraph()
        risk_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        risk_run = risk_para.add_run(f"RISK CLASS: {risk_class}")
        risk_run.font.size = Pt(16)
        risk_run.font.bold = True
        # Map hex to RGBColor
        _RISK_RGB = {
            "LOW":      RGBColor(0x27, 0xAE, 0x60),
            "MEDIUM":   RGBColor(0xF3, 0x9C, 0x12),
            "HIGH":     RGBColor(0xE6, 0x7E, 0x22),
            "CRITICAL": RGBColor(0xE7, 0x4C, 0x3C),
        }
        risk_run.font.color.rgb = _RISK_RGB.get(risk_class, RGBColor(0x95, 0xA5, 0xA6))

    doc.add_paragraph()

    reco = report.get("recommendation")
    if reco:
        reco_labels = {
            "ACCEPT":             "✅ ACCEPTANCE RECOMMENDED",
            "CONDITIONAL_ACCEPT": "⚠️ CONDITIONAL ACCEPTANCE",
            "REJECT_RECOMMENDED": "❌ REJECTION RECOMMENDED",
        }
        reco_para = doc.add_paragraph()
        reco_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        reco_run = reco_para.add_run(reco_labels.get(reco, reco))
        reco_run.font.size = Pt(13)
        reco_run.font.bold = True

    doc.add_paragraph("\n\n")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Analysis Date: {datetime.now(timezone.utc).strftime('%m/%d/%Y')}\n"
        f"Case Ref: {report.get('case_id', 'N/A')[:8]}...\n"
        f"Report ID: {report.get('report_id', 'N/A')[:8]}...\n"
        f"Produced by FinaCES V1.2"
    ).font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)


def _add_table_of_contents(doc) -> None:
    from docx.shared import Cm
    doc.add_heading("Table of Contents", level=1)
    for key, title in SECTION_TITLES.items():
        toc_para = doc.add_paragraph(title, style="List Number")
        toc_para.paragraph_format.left_indent = Cm(0.5)


def _add_all_sections(doc, report: dict) -> None:
    for key, title in SECTION_TITLES.items():
        content = report.get(key, "")
        doc.add_heading(title, level=1)
        if content:
            _add_markdown_content(doc, content)
        else:
            from docx.shared import RGBColor
            p = doc.add_paragraph("Section not available.")
            p.runs[0].font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)
        doc.add_paragraph()


def _add_markdown_content(doc, content: str) -> None:
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.strip() in ("---", "***", "___"):
            _add_horizontal_rule(doc)
        elif line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|---"):
            table_lines = [line]
            i += 2  # skip the separator
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_markdown_table(doc, table_lines)
            continue
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_rich_run(p, line[2:].strip())
        elif not line.strip():
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            _add_rich_run(p, line.strip())
        i += 1


def _add_rich_run(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _add_markdown_table(doc, lines: list) -> None:
    if not lines:
        return
    rows = [[c.strip() for c in line.strip().strip("|").split("|")] for line in lines]
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx]
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < n_cols:
                cell = row.cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                if row_idx == 0:
                    run = p.add_run(_strip_markdown(cell_data))
                    run.bold = True
                    _set_cell_background(cell, "D6EAF8")
                else:
                    _add_rich_run(p, cell_data)


def _add_horizontal_rule(doc) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_footer(doc, report: dict) -> None:
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run(
            f"FinaCES V1.2 — Confidential Report — "
            f"Report {report.get('report_id', 'N/A')[:8]} — "
            f"MCC/MCA Financial Evaluation Workbench"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)


def _add_page_break(doc) -> None:
    doc.add_page_break()


def _set_cell_background(cell, hex_color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tc_pr.append(shd)


# ════════════════════════════════════════════════════════════════
# SYNCHRONOUS GENERATION — PDF (executed via asyncio.to_thread)
# ════════════════════════════════════════════════════════════════

def _sync_generate_pdf(
    report:  dict,
    pdf_path: str,
    subdir:  str,
    ref:     str,
    ts:      str,
) -> str:
    """Generates the PDF file synchronously (called via to_thread).
    Returns the actual path of the created file (PDF or HTML fallback)."""
    html_content = _build_html_report(report)
    try:
        from weasyprint import HTML, CSS
        HTML(string=html_content).write_pdf(
            pdf_path,
            stylesheets=[CSS(string=_pdf_css())],
        )
        return pdf_path
    except ImportError:
        # Fallback: save as HTML if WeasyPrint not installed
        html_fallback = Path(subdir) / f"Note_MCC_{ref}_{ts}.html"
        html_fallback.write_text(html_content, encoding="utf-8")
        logger.warning("WeasyPrint not installed — HTML fallback generated: %s", html_fallback)
        return str(html_fallback)
    except Exception as exc:
        raise RuntimeError(f"PDF generation error: {exc}") from exc


def _build_html_report(report: dict) -> str:
    risk_class = _extract_risk_class(report)
    risk_color = RISK_COLORS_HEX.get(risk_class, "#95A5A6")

    reco = report.get("recommendation")
    reco_labels = {
        "ACCEPT":             "✅ Acceptance Recommended",
        "CONDITIONAL_ACCEPT": "⚠️ Conditional Acceptance",
        "REJECT_RECOMMENDED": "❌ Rejection Recommended",
    }
    reco_text = reco_labels.get(reco, "Not determined")
    today = datetime.now(timezone.utc).strftime("%m/%d/%Y")

    sections_html = ""
    for key, title in SECTION_TITLES.items():
        content = report.get(key, "")
        content_html = (
            _markdown_to_html(content) if content
            else "<p class='missing'>Section not available.</p>"
        )
        sections_html += (
            f'<section id="{key}">'
            f'<h2>{title}</h2>'
            f'{content_html}'
            f'</section>'
        )

    toc_items = "".join(
        f'<li><a href="#{k}">{t}</a></li>' for k, t in SECTION_TITLES.items()
    )

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>MCC-Grade Financial Analysis Report</title>
</head>
<body>
<div class="cover-page">
    <div class="cover-content">
        <h1 class="cover-title">FINANCIAL ANALYSIS REPORT</h1>
        <p class="cover-subtitle">Standard MCC/MCA — Fiduciary Grade</p>
        <div class="risk-badge" style="color:{risk_color}; border-color:{risk_color};">
            RISK CLASS: {risk_class}
        </div>
        <div class="reco-badge">{reco_text}</div>
        <div class="cover-meta">
            <p>Analysis Date: {today}</p>
            <p>Case Ref: {report.get('case_id', 'N/A')[:8]}...</p>
            <p>Report ID: {report.get('report_id', 'N/A')[:8]}...</p>
            <p class="confidential">CONFIDENTIAL — FinaCES V1.2</p>
        </div>
    </div>
</div>
<div class="toc">
    <h2>Table of Contents</h2>
    <ol>{toc_items}</ol>
</div>
<div class="report-body">{sections_html}</div>
<div class="doc-footer">
    <p>FinaCES V1.2 — Confidential Report —
    Report {report.get('report_id', 'N/A')[:8]} —
    MCC/MCA Financial Evaluation Workbench</p>
</div>
</body>
</html>"""


def _markdown_to_html(text: str) -> str:
    if not text:
        return ""
    lines = text.split("\n")
    result = []
    in_table = False
    in_list = False
    table_rows = []

    for line in lines:
        if line.startswith("### "):
            if in_list: result.append("</ul>"); in_list = False
            result.append(f"<h3>{_md_inline(line[4:])}</h3>")
        elif line.startswith("## "):
            if in_list: result.append("</ul>"); in_list = False
            result.append(f"<h3>{_md_inline(line[3:])}</h3>")
        elif line.strip() in ("---", "***"):
            if in_list: result.append("</ul>"); in_list = False
            result.append("<hr>")
        elif line.startswith("|"):
            if not in_table:
                in_table = True
                table_rows = []
            if re.match(r"^\|[\s\-:|]+\|", line):
                continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_rows.append(cells)
        else:
            if in_table:
                result.append(_table_to_html(table_rows))
                in_table = False
                table_rows = []
            if line.startswith("- ") or line.startswith("* "):
                if not in_list:
                    result.append("<ul>")
                    in_list = True
                result.append(f"<li>{_md_inline(line[2:])}</li>")
            else:
                if in_list:
                    result.append("</ul>")
                    in_list = False
                if not line.strip():
                    result.append("<br>")
                else:
                    result.append(f"<p>{_md_inline(line)}</p>")

    if in_table:
        result.append(_table_to_html(table_rows))
    if in_list:
        result.append("</ul>")
    return "\n".join(result)


def _table_to_html(rows: list) -> str:
    if not rows:
        return ""
    html = '<table class="data-table">'
    for i, row in enumerate(rows):
        html += "<tr>"
        tag = "th" if i == 0 else "td"
        for cell in row:
            html += f"<{tag}>{_md_inline(cell)}</{tag}>"
        html += "</tr>"
    html += "</table>"
    return html


def _md_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"\*(.+?)\*",     r"<em>\1</em>",         text)
    return text


def _pdf_css() -> str:
    return """
    * { box-sizing: border-box; margin: 0; padding: 0; }
    body { font-family: Calibri, Arial, sans-serif; font-size: 11pt; color: #2C3E50; line-height: 1.6; }
    .cover-page { page-break-after: always; display:flex; align-items: center; justify-content: center; min-height: 100vh; text-align: center; padding: 60px 40px; }
    .cover-title { font-size: 28pt; font-weight: bold; color: #1A527A; margin-bottom: 16px; }
    .cover-subtitle { font-size: 14pt; color: #2C3E50; margin-bottom: 40px; }
    .risk-badge { display: inline-block; font-size: 18pt; font-weight: bold; border: 3px solid; padding: 12px 32px; border-radius: 8px; margin-bottom: 24px; }
    .reco-badge { font-size: 13pt; font-weight: bold; margin-bottom: 40px; }
    .cover-meta p { color: #7F8C8D; font-size: 10pt; margin: 4px 0; }
    .confidential { font-weight: bold; color: #E74C3C !important; margin-top: 16px !important; }
    .toc { page-break-after: always; padding: 40px; }
    .toc h2 { font-size: 16pt; color: #1A527A; margin-bottom: 20px; }
    .toc ol li { margin: 8px 0; font-size: 11pt; }
    .toc a { color: #2C3E50; text-decoration: none; }
    section { padding: 30px 40px; page-break-before: always; }
    section h2 { font-size: 14pt; color: #1A527A; border-bottom: 2px solid #1A527A; padding-bottom: 8px; margin-bottom: 20px; }
    section h3 { font-size: 12pt; color: #2C3E50; margin: 20px 0 10px; }
    section p { margin-bottom: 10px; text-align: justify; }
    table.data-table { width: 100%; border-collapse: collapse; margin: 16px 0; font-size: 10pt; }
    table.data-table th { background-color: #D6EAF8; color: #1A527A; font-weight: bold; padding: 8px 10px; border: 1px solid #AED6F1; text-align: left; }
    table.data-table td { padding: 6px 10px; border: 1px solid #D6EAF8; }
    ul { margin: 10px 0 10px 24px; } ul li { margin-bottom: 6px; }
    hr { border: none; border-top: 1px solid #AAAAAA; margin: 16px 0; }
    strong { font-weight: bold; } em { font-style: italic; }
    .missing { color: #95A5A6; font-style: italic; }
    .doc-footer { text-align: center; font-size: 8pt; color: #95A5A6; border-top: 1px solid #EEE; padding: 12px 40px; margin-top: 40px; }
    @page { size: A4; margin: 2.5cm 2.5cm 3cm 3cm; @bottom-center { content: counter(page) " / " counter(pages); font-size: 9pt; color: #95A5A6; } }
    """


# ════════════════════════════════════════════════════════════════
# COMMON HELPERS
# ════════════════════════════════════════════════════════════════

def _extract_risk_class(report: dict) -> str:
    section_10 = report.get("section_10_scoring", "")
    for risk in ["CRITICAL", "HIGH", "MEDIUM", "LOW"]:
        if risk in (section_10 or ""):
            return risk
    return "N/A"


def _sanitize_filename(text: str) -> str:
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"\s+", "_", text.strip())
    return text[:40] or "rapport"


def _strip_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*",     r"\1", text)
    return text.strip()
